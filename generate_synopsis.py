import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_synopsis():
    doc = Document()

    # Define Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("BIN300: MINI PROJECT — 2025-26 Even Semester")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("---")

    # Title Section
    title_sec = doc.add_paragraph()
    title_sec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_sec.add_run("SYNOPSIS")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    run = p.add_run("Project Team No: ")
    run.bold = True
    p.add_run("BBIN6_15") # Placeholder

    # Team Table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Register No.'
    hdr_cells[1].text = 'Name'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Placeholders for names
    table.rows[1].cells[0].text = "126156127"
    table.rows[1].cells[1].text = "User Name 1"
    table.rows[2].cells[0].text = "126XXXXXX"
    table.rows[2].cells[1].text = "User Name 2"
    table.rows[3].cells[0].text = "126XXXXXX"
    table.rows[3].cells[1].text = "User Name 3"

    doc.add_paragraph("---")

    # Project Details
    p = doc.add_paragraph()
    run = p.add_run("Project Title: ")
    run.bold = True
    p.add_run("Integrative Multi-Omic Profiling of Gastric Cancer: Decoding the Microbiome–Host Transcriptome Axis and Causal Drivers of Disease Progression")

    p = doc.add_paragraph()
    run = p.add_run("Name of the Guide: ")
    run.bold = True
    p.add_run("Dr. B. Bioinformatics, Assistant Professor, SCBT, SASTRA Deemed to be University")

    doc.add_paragraph("---")

    # Abstract Section
    doc.add_heading('Abstract', level=2)
    abstract_text = (
        "Stomach Adenocarcinoma (STAD) remains a leading cause of global cancer mortality, with its pathogenesis deeply intertwined with gastric microbiome dysbiosis. "
        "This project presents a comprehensive multi-omic analytical framework to investigate the molecular crosstalk between the gastric microbiome and the host transcriptome. "
        "Utilizing data from TCGA-STAD, GTEx v10, and multiple GEO cohorts (GSE27342, GSE63089), we developed a robust pipeline optimized for the NVIDIA H200 DGX platform. "
        "Initial results confirmed a significant 'oralization' of the gastric microbiome, characterized by an enrichment of periodontal pathogens such as Fusobacterium and Streptococcus in tumor tissues. "
        "Differential gene expression analysis (DESeq2/Limma) identified over 4,000 dysregulated genes in tumors, predominantly involved in Epithelial-to-Mesenchymal Transition (EMT), Calcium Signaling, and Fanconi Anemia pathways. "
        "To bridge the gap between correlation and causation, we employed Bidirectional Mendelian Randomization (TwoSampleMR), identifying specific microbial taxa as potential causal drivers of gastric inflammation. "
        "Furthermore, modern network analysis (hdWGCNA and scLink) revealed highly connected gene modules associated with Lauren subtypes, highlighting a more immunosuppressive microenvironment in the Diffuse subtype. "
        "Immune infiltration profiling (TIMER/xCell) further corroborated these findings, demonstrating a strong correlation between EMT activation and CD8+ T cell exclusion. "
        "Our findings provide a systems-level understanding of gastric cancer progression and nominate novel microbiome-modulated pathways as potential therapeutic targets."
    )
    
    p = doc.add_paragraph(abstract_text)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15

    doc.add_paragraph("---")

    # Specific Contribution
    doc.add_heading('Specific Contribution', level=2)
    contributions = [
        "Development and optimization of a multi-omic integration pipeline for high-performance computing (HPC) environments.",
        "Harmonization of multi-cohort transcriptomic data (TCGA, GTEx, GEO) using ComBat to eliminate technical batch effects.",
        "Implementation of a stringent MaAsLin2 modeling framework to identify robust microbiome-transcriptome associations.",
        "Execution of Mendelian Randomization analysis to prioritize causal microbial taxa in gastric cancer risk.",
        "Profiling of the tumor microenvironment (TME) through automated immune deconvolution and pathway correlation modules."
    ]
    for item in contributions:
        doc.add_paragraph(item, style='List Bullet')

    # Specific Learning
    doc.add_heading('Specific Learning', level=2)
    learnings = [
        "Advanced bioinformatics techniques for processing and normalizing large-scale bulk RNA-seq and 16S microbiome datasets.",
        "Theoretical and practical application of causal inference models (Mendelian Randomization) in clinical genomics.",
        "Network-based analysis methods (WGCNA, scLink) for identifying co-expression patterns and biological modules.",
        "Parallel programming and resource management on NVIDIA H200 DGX hardware for large-scale multi-omics integration.",
        "Understanding the immunological landscape of gastric cancer Lauren subtypes (Diffuse vs. Intestinal)."
    ]
    for item in learnings:
        doc.add_paragraph(item, style='List Bullet')

    # Technical Limitations
    doc.add_heading('Technical Limitations & Ethical Challenges Faced', level=2)
    limitations = [
        "Computational constraints related to R's parallel socket limits when scaling to 200+ cores on the DGX server.",
        "Instability and rate-limiting of public BioMart servers, necessitating the implementation of local annotation fallbacks.",
        "Ethical considerations regarding the use of de-identified patient data from the Cancer Genome Atlas (TCGA) and dbGaP.",
        "Technical challenges in aligning disparate multi-omic cohorts due to varying sample naming conventions and metadata formats."
    ]
    for item in limitations:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("---")

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run("Keywords: Gastric Cancer, Multi-omics Integration, Microbiome Dysbiosis, Mendelian Randomization, Tumor Microenvironment")
    run.italic = True
    run.font.size = Pt(10)

    doc.add_paragraph("---")

    # Signatures
    p = doc.add_paragraph()
    p.add_run("Signature of the Student")
    p.add_run(" " * 40)
    p.add_run("Signature of Guide with Date")

    # Save
    filename = "Gastric_Cancer_MultiOmics_Synopsis.docx"
    doc.save(filename)
    print(f"Synopsis successfully saved to {filename}")

if __name__ == "__main__":
    generate_synopsis()
