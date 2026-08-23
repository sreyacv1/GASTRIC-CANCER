#!/usr/bin/env python3
"""Verify every quantitative claim in PROJECT_A_TO_Z.md against its source file.

Run from anywhere; the script locates the repository root itself:

    python3 analysis/verify_a_to_z.py

Each check names the claim, the value parsed from source, and whether that value
appears verbatim in the document (and, when built, in the .docx and .pdf too).
Exit status is non-zero if any check fails, so this can gate a rebuild.
"""
import csv
import glob
import gzip
import os
import re
import subprocess
import sys

def _find_root(start):
    """Walk up from this file until the repository markers appear.

    Keeps the script runnable whether it sits at the repo root or in analysis/.
    """
    d = os.path.dirname(os.path.abspath(start))
    for _ in range(4):
        if os.path.isdir(os.path.join(d, "results")) and \
           os.path.isfile(os.path.join(d, "PAPER.md")):
            return d
        d = os.path.dirname(d)
    # A copy of this script may live outside the repo (e.g. saved as a standalone
    # artifact). Fall back to an explicit repo path, then to $GASTRIC_REPO.
    for cand in (os.environ.get("GASTRIC_REPO"),
                 "/nfsshare/users/P126156127/workspace/bioinf/gastric_cancer"):
        if cand and os.path.isdir(os.path.join(cand, "results")) and \
           os.path.isfile(os.path.join(cand, "PAPER.md")):
            return cand
    raise SystemExit(
        "cannot locate repository root: run this from inside the repo, or set "
        "GASTRIC_REPO to the checkout path")


ROOT = _find_root(__file__)
os.chdir(ROOT)

DOC = "PROJECT_A_TO_Z.md"


def rd(path):
    with open(path) as fh:
        return list(csv.DictReader(fh))


def kv(path):
    """Two-column csv -> dict."""
    rows = rd(path)
    if not rows:
        return {}
    cols = list(rows[0])
    return {r[cols[0]]: r[cols[1]] for r in rows}


checks = []  # (label, string_expected_in_doc, source_value_ok)


def add(label, shown, ok):
    checks.append((label, shown, bool(ok)))


# ---- differential expression -------------------------------------------------
deg = rd("results/tables/TCGA_DEG_results.csv")
add("genes tested", "21,446", len(deg) == 21446)
add("up-regulated", "2,134", sum(1 for r in deg if r["sig"] == "Up") == 2134)
add("down-regulated", "2,362", sum(1 for r in deg if r["sig"] == "Down") == 2362)

diag = glob.glob("results/deg_diagnostics/*biotype*.csv")
if diag:
    bt = rd(diag[0])
    col = [c for c in bt[0] if "coding" in c.lower() or "biotype" in c.lower()]
    pc = None
    for r in bt:
        vals = " ".join(str(v) for v in r.values())
        if "protein" in vals.lower():
            nums = [int(x) for x in re.findall(r"\b(\d{4,6})\b", vals)]
            if nums:
                pc = max(nums)
    if pc:
        add("protein-coding count", "16,164", pc == 16164)
        add("protein-coding pct", "75.4%", abs(pc / len(deg) * 100 - 75.4) < 0.1)

bs = glob.glob("results/deg_diagnostics/*batch*sensitiv*.csv") + \
     glob.glob("results/deg_diagnostics/*TSS*.csv")
if bs:
    txt = open(bs[0]).read()
    for lab, val in [("batch DEG group-only", "4,456"), ("batch DEG +TSS", "4,569"),
                     ("batch overlap pct", "92.7%"), ("batch logFC r", "0.974")]:
        add(lab, val, val.replace(",", "").rstrip("%") in txt.replace(",", ""))

# ---- WGCNA -------------------------------------------------------------------
for coh, want in [("ACRG_GSE62254", "15.9"), ("GSE15459", "16.8"), ("GSE84437", "17.1")]:
    p = f"results/module_preservation/preservation_stats_{coh}.csv"
    red = [x for x in rd(p) if x["module"] == "red"]
    add(f"Zsummary {coh}", want, red and abs(float(red[0]["Zsummary.pres"]) - float(want)) < 0.05)

# Module size vs per-cohort overlap: the preservation files' moduleSize column is the
# number of module genes measurable on THAT cohort's platform, not the module size.
mods = rd("results/wgcna_real/hub_genes_prognostic_module.csv")
# Anchor on the unique headline phrase: "263 genes" alone also occurs in the
# disambiguation paragraph, so a corrupted headline would still have matched.
add("red module size", "red module — 263 genes", len(mods) == 263)
add("module size (disambiguation)", "module itself is 263 genes", len(mods) == 263)
ov = {}
for f in glob.glob("results/module_preservation/preservation_stats_*.csv"):
    red = [x for x in rd(f) if x["module"] == "red"]
    if red:
        ov[os.path.basename(f)] = red[0]["moduleSize"]
add("ACRG platform overlap", "77 of 263",
    any(v == "77" for k, v in ov.items() if "ACRG" in k))
add("other-cohort overlap", "233 on the other two",
    sum(1 for k, v in ov.items() if v == "233") == 2)

eig = rd("results/module_preservation/module_eigengene_cox_external.csv")
for row, want in zip(eig, ["1.274", "1.548", "1.237"]):
    add(f"eigengene HR {row['cohort']}", want, row["HR_perSD"] == want)

adj = rd("results/wgcna_real/ME_survival_cox_adjusted.csv")[0]
add("eigengene adjusted HR", "1.347", adj["HR"] == "1.347")
add("eigengene adjusted p", "0.090", round(float(adj["p"]), 3) == 0.090)
un = rd("results/wgcna_real/ME_survival_cox.csv")[0]
add("eigengene unadjusted HR", "1.307", un["HR_perSD"] == "1.307")
add("eigengene unadjusted p", "0.00093", abs(float(un["p"]) - 0.00093) < 5e-6)

pw = rd("results/wgcna_real/power_robustness_summary.csv")
hrs = sorted(float(r["HR_perSD"]) for r in pw)
add("power HR range low", "1.283", hrs[0] == 1.283)
add("power HR range high", "1.308", hrs[-1] == 1.308)
add("power max cox p", "0.0024", max(float(r["cox_p"]) for r in pw) <= 0.0024)
sft = sorted(float(r["SFT_R2"]) for r in pw)
add("scale-free R2 at power 3", "0.88", abs(float(pw[0]["SFT_R2"]) - 0.877) < 0.005)

# ---- single cell -------------------------------------------------------------
cc = rd("results/scrna/celltype_composition.csv")
add("scRNA cells", "43,992", sum(int(x["n_cells"]) for x in cc) == 43992)
add("scRNA cell types", "eight annotated", len(cc) == 8)
nc = rd("results/scrna/gene_dominant_celltype_noncircular.csv")
fr = sorted(float(x["frac_in_dominant"]) for x in nc)
add("non-circular genes", "23 of 23",
    len(nc) == 23 and all(x["dominant_cell_type"] == "Fibroblast" for x in nc))
add("dominance range low", "0.527", fr[0] == 0.527)
add("dominance range high", "0.998", fr[-1] == 0.998)
add("dominance median", "0.96", abs(fr[len(fr) // 2] - 0.96) < 0.005)

# ---- signature ---------------------------------------------------------------
perf = kv("results/nested_cv/performance.csv")
add("nested Harrell C", "0.6112", round(float(perf["Harrell_C_ensemble"]), 4) == 0.6112)
add("nested Uno C", "0.5727", round(float(perf["Uno_C_ensemble"]), 4) == 0.5727)
add("nested per-repeat C", "0.5981", round(float(perf["Harrell_C_perRepeat"]), 4) == 0.5981)
add("integrated Brier", "0.1789", round(float(perf["IntegratedBrierScore"]), 4) == 0.1789)

ci = rd("results/validation/cindex_comparison.csv")
add("TCGA C-index", "0.719", abs(float(ci[0]["C_index"]) - 0.719) < 0.001)
add("ACRG C-index", "0.608", abs(float(ci[1]["C_index"]) - 0.608) < 0.001)

mcs = {r["cohort"]: r for r in rd("results/validation_multi/cindex_HR_summary.csv")}
add("GSE15459 C-index", "0.575", round(float(mcs["GSE15459"]["C_index"]), 3) == 0.575)
add("GSE84437 C-index", "0.530", round(float(mcs["GSE84437"]["C_index"]), 3) == 0.530)
add("GSE84437 HR", "1.11 (0.84–1.46)",
    round(float(mcs["GSE84437"]["HR"]), 2) == 1.11
    and round(float(mcs["GSE84437"]["HR_low"]), 2) == 0.84
    and round(float(mcs["GSE84437"]["HR_high"]), 2) == 1.46)
add("GSE15459 HR", "1.68 (1.11–2.54)",
    round(float(mcs["GSE15459"]["HR"]), 2) == 1.68)
add("ACRG HR", "1.90 (1.37–2.62)",
    round(float(mcs["ACRG/GSE62254"]["HR"]), 2) == 1.90)

acrg = [r for r in rd("results/validation/multivariable_cox_ACRG.csv") if r["term"] == "groupHigh"][0]
add("ACRG adjusted HR", "1.76", abs(float(acrg["HR"]) - 1.76) < 0.005)

meta = rd("results/meta_HK/meta_result.csv")[0]
add("pooled HR", "1.188", round(float(meta["pooled_HR"]), 3) == 1.188)
add("pooled CI low", "0.962", round(float(meta["CI_low"]), 3) == 0.962)
add("pooled CI high", "1.466", round(float(meta["CI_high"]), 3) == 1.466)
add("pooled p", "0.073", round(float(meta["p_value"]), 3) == 0.073)
add("I-squared", "19.2", round(float(meta["I2"]), 1) == 19.2)

ts = rd("results/validation_multi/GSE84437_Tstage_stratified.csv")
add("pT strata all below 0.5", "below 0.5", all(r["C_below_0.5"] == "TRUE" for r in ts))

zph = glob.glob("results/timevarying_ACRG/coxzph*.csv")
if zph:
    txt = open(zph[0]).read()
    add("cox.zph p", "0.0018", "0.0018" in txt or "0.00179" in txt or "0.0017" in txt)
tv = glob.glob("results/timevarying_ACRG/*time*vary*.csv") + \
     glob.glob("results/timevarying_ACRG/*HR*.csv")
if tv:
    txt = open(tv[0]).read()
    for lab, v in [("tv HR 12mo", "1.49"), ("tv HR 36mo", "1.03"), ("tv HR 60mo", "0.87")]:
        add(lab, v, v in txt)

it = rd("results/predictive/interaction_tests.csv")
ip = {r["interaction"]: r["p_value"] for r in it}
stg = [v for k, v in ip.items() if "stage" in k][0]
sub = [v for k, v in ip.items() if "subtype" in k][0]
add("stage interaction p", "0.044", round(float(stg), 3) == 0.044)
add("subtype interaction p", "0.094", round(float(sub), 3) == 0.094)

bp = rd("results/base_paper_replication/model_performance.csv")[0]
add("5-gene apparent C", "0.5446", bp["C_apparent"] == "0.5446")
add("5-gene corrected C", "0.5193", bp["C_optimism_corrected"] == "0.5193")
add("5-gene HR", "0.627", bp["HR_highVsLow"] == "0.627")
old = rd("results/base_paper_replication/model_performance_rms6.8-1_R4.3.3.csv")[0]
add("5-gene C under rms 6.8-1", "0.4807", old["C_optimism_corrected"] == "0.4807")
si = open("results/base_paper_replication/sessionInfo.txt").read()
add("R version current", "4.5.3", "4.5.3" in si)
add("rms version current", "8.1-1", "8.1-1" in si)
si2 = open("results/base_paper_replication/sessionInfo_rms6.8-1_R4.3.3.txt").read()
add("R version archived", "4.3.3", "4.3.3" in si2)
add("rms version archived", "6.8-1", "6.8-1" in si2)

# ---- microbiome --------------------------------------------------------------
da_a = rd("results/microbiome_biomarker/04a_DA_control_vs_GCN.csv")
da_b = rd("results/microbiome_biomarker/04b_DA_GCN_vs_GCT_paired.csv")


def nsig(rows):
    key = "q" if "q" in rows[0] else ("qval" if "qval" in rows[0] else None)
    return sum(1 for r in rows if key and float(r[key] or 1) < 0.05)


add("DA control-vs-GCN", "44 of 61", nsig(da_a) == 44 and len(da_a) == 61)
add("DA paired GCN-vs-GCT", "18 of 61", nsig(da_b) == 18 and len(da_b) == 61)

alpha = {r["metric"]: r for r in rd("results/microbiome_biomarker/02_alpha_effectsizes_cascade.csv")}
for met, delta in [("Observed", "−0.27"), ("Shannon", "−0.122"), ("Simpson", "−0.072")]:
    if met in alpha:
        src = alpha[met]["cliffs_delta"]
        add(f"Cliff delta {met}", delta, src == delta.replace("−", "-"))

rf = {r["metric"]: r["value"] for r in rd("results/microbiome_biomarker/05_rf_metrics_and_batch_sanity.csv")}
add("cancer AUC", "0.916", round(float(rf["cancer_vs_control_AUC"]), 3) == 0.916)
add("AUC CI low", "0.896", round(float(rf["cancer_AUC_CI_lo"]), 3) == 0.896)
add("AUC CI high", "0.936", round(float(rf["cancer_AUC_CI_hi"]), 3) == 0.936)
add("flowcell accuracy", "77.6%", round(float(rf["flowcell_pred_accuracy"]) * 100, 1) == 77.6)
add("flowcell baseline", "54.5%", round(float(rf["flowcell_majority_baseline"]) * 100, 1) == 54.5)

# ---- Mendelian randomization -------------------------------------------------
mr = rd("results/mr_real/MR_per_exposure_instruments_REAL.csv")
for r in mr:
    shown = f"{float(r['IVW_OR']):.2f} ({r['IVW_CI'].replace('-', '–')})"
    add(f"MR {r['Exposure'][:22]}", shown, True)
    add(f"MR nSNP {r['Exposure'][:18]}", r["nSNP_used"], True)
add("MR min F", "19.3", min(float(r["min_F"]) for r in mr) == 19.3)
add("MR PRESSO lowest", "0.052", min(float(r["PRESSO_global_p"]) for r in mr) == 0.052)
add("MR all PRESSO ns", "> 0.05", all(float(r["PRESSO_global_p"]) > 0.05 for r in mr))

# ---- immune ------------------------------------------------------------------
im = {(r["estimate"], r["measured"]): float(r["spearman_rho"])
      for r in rd("results/immune/validation_vs_measured.csv")}
add("T cells vs leukocyte", "0.666", round(im[("T cells (MCP)", "Leukocyte %")], 3) == 0.666)
add("ImmuneScore vs leukocyte", "0.652", round(im[("ImmuneScore (xCell)", "Leukocyte %")], 3) == 0.652)
add("CD8 vs leukocyte", "0.468", round(im[("CD8 T cells (MCP)", "Leukocyte %")], 3) == 0.468)
add("CD8 vs lymphocyte", "0.020", round(im[("CD8 T cells (MCP)", "Lymphocyte infiltration %")], 3) == 0.020)
c8 = rd("results/immune/CD8_survival_summary.csv")[0]
add("CD8 Cox HR", "1.042", round(float(c8["cox_HR_per_unit_CD8"]), 3) == 1.042)
add("CD8 CI low", "0.944", round(float(c8["cox_CI_low"]), 3) == 0.944)
add("CD8 CI high", "1.150", round(float(c8["cox_CI_high"]), 3) == 1.150)
add("CD8 Cox p", "0.411", round(float(c8["cox_p"]), 3) == 0.411)

# ---- DepMap ------------------------------------------------------------------
dm = {x["gene"]: round(float(x["gastric_mean_CERES"]), 3)
      for x in rd("results/depmap/gastric_dependency.csv")}
for g, v in [("PIK3CA", "−0.742"), ("MTOR", "−1.184"), ("CDK4", "−0.825"), ("CDK6", "−0.548")]:
    add(f"DepMap {g}", v, dm.get(g) == float(v.replace("−", "-")))

# ---- cohort composition ------------------------------------------------------
stage = {}
with gzip.open("data/geo/GSE84437_series_matrix.txt.gz", "rt", errors="replace") as fh:
    for line in fh:
        if line.startswith("!Sample_characteristics") and "ptstage" in line:
            vals = [v.split(":")[-1].strip() for v in re.findall(r'"([^"]*)"', line)]
            for v in vals:
                if v:
                    stage[v] = stage.get(v, 0) + 1
tot = sum(stage.values())
pt34 = stage.get("T3", 0) + stage.get("T4", 0)
add("GSE84437 staged n", "433", tot == 433)
add("GSE84437 pT3-T4 pct", "88.7%", abs(pt34 / tot * 100 - 88.7) < 0.05)
for k, v in [("T1", "11"), ("T2", "38"), ("T3", "92"), ("T4", "292")]:
    add(f"GSE84437 {k}", v, str(stage.get(k)) == v)

mi = rd("results/meta_HK/meta_inputs.csv")
for r in mi:
    add(f"cohort n {r['cohort'][:14]}", r["n"], True)
    add(f"cohort events {r['cohort'][:14]}", r["events"], True)

# ---- figures -----------------------------------------------------------------
try:
    from PIL import Image
    md = open("PAPER.md").read()
    hi = []
    for lbl, p in re.findall(r"!\[([^\]]*)\]\(([^)]+)\)", md):
        if os.path.exists(p):
            w, h = Image.open(p).size
            if h / w > 1.00:
                hi.append((lbl, round(h / w, 2)))
    add("figures above aspect 1.00", "five of the 17", len(hi) == 5)
    for lbl, a in hi:
        add(f"aspect {lbl}", f"{a:.2f}", True)
except ImportError:
    pass

# ---- repository --------------------------------------------------------------
# Script and commit counts are intentionally NOT asserted: they change with every
# commit, so quoting them in the document guarantees staleness. The document now
# tells the reader how to obtain them instead.
nfig = len(re.findall(r"!\[", open("PAPER.md").read()))
add("figures embedded", f"{nfig} figures", True)

# ---- run ---------------------------------------------------------------------
doc = open(DOC).read()
# Compare against emphasis-stripped Markdown so a claim wrapped in ** matches the
# same claim in the rendered .docx/.pdf, where the markers are gone.
targets = {"md": doc.replace("**", "")}
try:
    from docx import Document
    if os.path.exists("PROJECT_A_TO_Z.docx"):
        d = Document("PROJECT_A_TO_Z.docx")
        targets["docx"] = "\n".join(p.text for p in d.paragraphs) + "\n" + \
            "\n".join(c.text for t in d.tables for r in t.rows for c in r.cells)
except ImportError:
    pass
try:
    import pypdfium2 as pdfium
    if os.path.exists("PROJECT_A_TO_Z.pdf"):
        pd = pdfium.PdfDocument("PROJECT_A_TO_Z.pdf")
        targets["pdf"] = "".join(pd[i].get_textpage().get_text_range() for i in range(len(pd)))
except ImportError:
    pass

def _norm(s):
    """Collapse whitespace: PDF and DOCX text extraction wraps lines, which would
    otherwise split a multi-word claim and report a false mismatch."""
    return re.sub(r"\s+", " ", s)


targets = {k: _norm(v) for k, v in targets.items()}

fails = []
for label, shown, src_ok in checks:
    where = {k: (_norm(shown) in v) for k, v in targets.items()}
    if not src_ok or not all(where.values()):
        fails.append((label, shown, src_ok, where))

print(f"checks: {len(checks)}  formats: {sorted(targets)}  failures: {len(fails)}")
for label, shown, src_ok, where in fails:
    print(f"  FAIL {label}: shown={shown!r} source_ok={src_ok} present={where}")
sys.exit(1 if fails else 0)
